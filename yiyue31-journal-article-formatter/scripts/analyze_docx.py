#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
分析 docx 文档的结构和格式信息。
用法: python analyze_docx.py <文件名.docx>
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

def analyze(filepath):
    doc = Document(filepath)

    # 1. 页面设置
    print('=== 页面设置 ===')
    for i, section in enumerate(doc.sections):
        ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        cols = section._sectPr.findall(f'{ns}cols')
        col_info = '1栏(默认)'
        if cols:
            c = cols[0]
            num = c.get(f'{ns}num', '1')
            space = c.get(f'{ns}space', '720')
            col_info = f'{num}栏, 间距={space}'
        print(f'  Section {i}: {section.page_width.cm:.2f}x{section.page_height.cm:.2f}cm')
        print(f'  边距: 上{section.top_margin.cm:.2f} 下{section.bottom_margin.cm:.2f} 左{section.left_margin.cm:.2f} 右{section.right_margin.cm:.2f}')
        print(f'  栏: {col_info}')
        print(f'  页眉距离: {section.header_distance.cm:.2f}cm, 页脚距离: {section.footer_distance.cm:.2f}cm')
        print()

    # 2. 段落详情
    print('=== 段落详情 ===')
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else 'None'
        alignment = para.alignment
        pf = para.paragraph_format
        first_line = pf.first_line_indent

        first_run = para.runs[0] if para.runs else None
        if first_run:
            font = first_run.font.name
            size = round(first_run.font.size / 12700, 1) if first_run.font.size else None
            bold = first_run.bold
        else:
            font = size = bold = None

        display = text[:80] + '...' if len(text) > 80 else text
        indent_info = f' 首行缩进={first_line}' if first_line else ''
        print(f'[{i}] 字体={font} 大小={size}pt 粗体={bold} 对齐={alignment}{indent_info}')
        print(f'     {display}')

    # 3. 页眉页脚
    print()
    print('=== 页眉页脚 ===')
    for i, section in enumerate(doc.sections):
        header = section.header
        footer = section.footer
        h_text = ' | '.join(p.text.strip() for p in header.paragraphs if p.text.strip()) or '无'
        f_text = ' | '.join(p.text.strip() for p in footer.paragraphs if p.text.strip()) or '无'
        print(f'  Section {i}: 页眉="{h_text}"')
        print(f'  Section {i}: 页脚="{f_text}"')

    # 4. 表格和图片
    print()
    print(f'=== 其他 ===')
    print(f'  表格数: {len(doc.tables)}')
    rels = doc.part.rels
    img_count = sum(1 for rel in rels.values() if 'image' in rel.reltype)
    print(f'  图片数: {img_count}')

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('用法: python analyze_docx.py <文件名.docx>')
        sys.exit(1)
    analyze(sys.argv[1])
