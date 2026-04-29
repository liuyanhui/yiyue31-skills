#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
验证生成的期刊格式 docx 文件是否符合规范。
用法: python verify_output.py <文件名.docx>
"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

def verify(filepath, expected_cols=None):
    if not os.path.exists(filepath):
        print(f'文件不存在: {filepath}')
        sys.exit(1)

    doc = Document(filepath)
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    errors = []
    warnings = []

    # 默认栏数期望：首页单栏、正文双栏、参考文献单栏、引用格式单栏
    # 可通过命令行参数覆盖，如：python verify_output.py 文件.docx 1 2 1 1
    if expected_cols is None:
        expected_cols = [1, 2, 1, 1]

    print(f'文件大小: {os.path.getsize(filepath) / 1024:.1f} KB')
    print(f'Section 数量: {len(doc.sections)}')
    print(f'期望栏数: {expected_cols}')
    print()

    # 检查 Section 结构
    for i, section in enumerate(doc.sections):
        cols = section._sectPr.findall(f'{ns}cols')
        col_num = int(cols[0].get(f'{ns}num', '1')) if cols else 1
        expected = expected_cols[i] if i < len(expected_cols) else 1
        status = 'OK' if col_num == expected else 'FAIL'
        if status == 'FAIL':
            errors.append(f'Section {i}: 预期 {expected} 栏，实际 {col_num} 栏')
        print(f'  Section {i}: {col_num}栏 ... {status}')

    print()

    # 检查首页段落顺序
    print('=== 首页段落检查 ===')
    section_breaks = []
    for i, para in enumerate(doc.paragraphs):
        pPr = para._element.find(f'{ns}pPr')
        if pPr is not None:
            sectPr = pPr.find(f'{ns}sectPr')
            if sectPr is not None:
                section_breaks.append(i)

    end = section_breaks[0] + 1 if section_breaks else len(doc.paragraphs)
    first_page_texts = []
    for i in range(end):
        text = doc.paragraphs[i].text.strip()
        if text:
            first_page_texts.append((i, text))

    # 检查关键元素
    all_first_page = ' '.join(t for _, t in first_page_texts)
    checks = [
        ('摘要', '摘' in all_first_page or '摘  要' in all_first_page),
        ('关键词', '关  键  词' in all_first_page or '关键词' in all_first_page),
        ('Abstract', 'Abstract' in all_first_page),
        ('Key words', 'Key words' in all_first_page),
        ('中图分类号', '中图分类号' in all_first_page),
        ('参考文献', any('参考文献' in doc.paragraphs[i].text for i in range(len(doc.paragraphs)))),
    ]

    for name, ok in checks:
        status = 'OK' if ok else 'MISSING'
        if not ok:
            errors.append(f'首页缺少: {name}')
        print(f'  {name} ... {status}')

    print()

    # 检查上标引用
    print('=== 上标引用检查 ===')
    superscript_count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.superscript and run.text.strip().startswith('['):
                superscript_count += 1
    if superscript_count > 0:
        print(f'  找到 {superscript_count} 个上标引用 ... OK')
    else:
        warnings.append('未找到上标引用标注，正文中可能缺少 [n] 上标格式')
        print(f'  未找到上标引用 ... WARN')

    print()

    # 检查参考文献双语
    print('=== 参考文献双语检查 ===')
    ref_section = False
    ref_lines = 0
    has_in_chinese = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if '参考文献' in text and not ref_section:
            ref_section = True
            continue
        if ref_section and text:
            ref_lines += 1
            if '(in Chinese)' in text:
                has_in_chinese = True

    if ref_lines > 0:
        print(f'  参考文献条数: {ref_lines} ... OK')
    else:
        errors.append('参考文献内容为空')
        print(f'  参考文献内容为空 ... FAIL')

    if has_in_chinese:
        print(f'  双语格式 (in Chinese) ... OK')
    else:
        warnings.append('参考文献中未找到 "(in Chinese)" 标记，可能缺少英文翻译')
        print(f'  双语格式 (in Chinese) ... WARN')

    print()

    # 汇总
    print('=== 验证结果 ===')
    if errors:
        print(f'  错误 ({len(errors)}):')
        for e in errors:
            print(f'    - {e}')
    if warnings:
        print(f'  警告 ({len(warnings)}):')
        for w in warnings:
            print(f'    - {w}')
    if not errors and not warnings:
        print('  全部通过!')

    return len(errors) == 0

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('用法: python verify_output.py <文件名.docx> [栏数1 栏数2 ...]')
        print('  默认栏数: 1 2 1 1 (首页单栏、正文双栏、参考文献单栏、引用格式单栏)')
        print('  示例: python verify_output.py output.docx 1 2 1')
        sys.exit(1)
    cols = None
    if len(sys.argv) > 2:
        cols = [int(x) for x in sys.argv[2:]]
    ok = verify(sys.argv[1], expected_cols=cols)
    sys.exit(0 if ok else 1)
